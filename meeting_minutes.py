import streamlit as st
import tempfile
import os
from streamlit_mic_recorder import mic_recorder, speech_to_text
import openai
from docx import Document
import io  


def save_audio_to_tempfile(audio_bytes):
    # Check if 'bytes' key exists in audio_bytes
    if 'bytes' in audio_bytes:
        # Saving the recorded audio to a temporary file
        with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as f:
            f.write(audio_bytes['bytes'])
            audio_file_path = f.name
        return audio_file_path
    else:
        raise ValueError("audio_bytes does not contain 'bytes' key")

def transcribe_audio(audio_file_path):
    with open(audio_file_path, 'rb') as audio_file:
        response = openai.Audio.transcribe(model="whisper-1", file=audio_file)
    return response['text']

def meeting_minutes(transcription):
    abstract_summary = abstract_summary_extraction(transcription)
    key_points = key_points_extraction(transcription)
    action_items = action_item_extraction(transcription)
    sentiment = sentiment_analysis(transcription)
    return {
        'abstract_summary': abstract_summary,
        'key_points': key_points,
        'action_items': action_items,
        'sentiment': sentiment
    }

def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a skilled AI assistant for analyzing meetings and extracting action items. Given the following text of a meeting transcript, please identify and list out any action items that were discussed or assigned. Action items should be specific tasks or action points that need to be completed by someone. Your response should be a clear, bulleted list of these action items."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanation in 10-20 words for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response.choices[0].message['content']

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()
    doc.save(filename)

def generate_docx_bytes(minutes):
    doc = Document()
    for key, value in minutes.items():
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        doc.add_paragraph()

    # Create a BytesIO stream to store the Word document
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)  

    return docx_bytes

def app():
    st.title("Meeting Minutes Generator")

    # Get the OpenAI API key from the user
    api_key = st.text_input("Enter your OpenAI API Key", type="password", key="openai_api_key_input")
    openai.api_key = api_key


    # Record audio using the microphone
    st.write("Record the meeting:")
    audio_bytes = mic_recorder(start_prompt="Start Recording", stop_prompt="Stop Recording", key='recorder')

    if audio_bytes:
        try:
            # Save the audio to a temporary file
            audio_file_path = save_audio_to_tempfile(audio_bytes)
            
            # Transcribe the audio
            transcription_text = transcribe_audio(audio_file_path)
            
            # Display the transcription text
            st.subheader("Transcription")
            st.write(transcription_text)
            
            # Generate meeting minutes
            minutes = meeting_minutes(transcription_text)
            
            # Display the meeting minutes
            st.header("Meeting Minutes")

            # Abstract Summary
            st.subheader("Abstract Summary")
            st.write(minutes['abstract_summary'])

            # Key Points
            st.subheader("Key Points")
            for point in minutes['key_points'].split('\n'):
                st.write(f"- {point.strip()}")

            # Action Items
            st.subheader("Action Items")
            for item in minutes['action_items'].split('\n'):
                st.write(f"- {item.strip()}")

            # Sentiment
            st.subheader("Sentiment of Meeting")
            st.write(minutes['sentiment'])

            # Generate the Word document as a byte stream
            docx_bytes = generate_docx_bytes(minutes)

            # Download button
            st.download_button(
                label="Download Meeting Minutes",
                data=docx_bytes.getvalue(),
                file_name="Meeting_Minutes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        except Exception as e:
            # Print any error that occurs
            st.error(f"An error occurred: {e}")

if __name__ == "__main__":
    app()
